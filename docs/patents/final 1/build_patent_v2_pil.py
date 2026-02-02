import argparse, importlib.util, re, zipfile
from pathlib import Path

from reportlab.lib.pagesizes import LETTER
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.graphics import renderPDF, renderPM

# Optional docx import (skip if missing)
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAVE_DOCX = True
except Exception:
    HAVE_DOCX = False

def log(m): print(f"[v2-pil] {m}")

def load_fig_module(p: Path):
    spec = importlib.util.spec_from_file_location("nbmf_figs_v2", str(p))
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)  # type: ignore
    funcs = getattr(mod, "FIG_FUNCS", [])
    caps = getattr(mod, "FIG_CAPTIONS", [])
    details = getattr(mod, "FIG_DETAILS", [])
    if not funcs or not caps:
        raise RuntimeError("FIG_FUNCS or FIG_CAPTIONS missing")
    return funcs, caps, details

def parse_claims(md: str):
    md = md.replace("\r\n","\n").replace("\r","\n")
    claims = []
    heads = list(re.finditer(r"^###\s*Claim\s*(\d+)\s*\((.*?)\)?:?\s*(.*)$", md, flags=re.M))
    if heads:
        for i, m in enumerate(heads):
            num = int(m.group(1))
            dep = (m.group(2) or "").strip()
            title = m.group(3).strip()
            start = m.end()
            end = heads[i+1].start() if i+1<len(heads) else len(md)
            body = re.sub(r"\s+", " ", md[start:end].strip())
            suffix = f" (depends on {dep})" if dep else ""
            claims.append((num, f"{title}{suffix}", body))
        return claims
    for ln in md.splitlines():
        m = re.match(r"^\s*(\d+)\.\s*(.+)", ln.strip())
        if m:
            claims.append((int(m.group(1)), "", m.group(2)))
    return claims

def claims_from_zip(zp: Path):
    with zipfile.ZipFile(str(zp), "r") as z:
        for name in z.namelist():
            if name.lower().endswith("claims_draft.md"):
                return parse_claims(z.read(name).decode("utf-8", errors="ignore"))
        for name in z.namelist():
            if name.lower().endswith(".md"):
                c = parse_claims(z.read(name).decode("utf-8", errors="ignore"))
                if c: return c
    return []

def build_pdf(out_pdf: Path, funcs, caps, details, claims, compact=False):
    styles = getSampleStyleSheet()
    S_TITLE = styles['Title']
    S_H1 = ParagraphStyle(name='H1', parent=styles['Heading1'], spaceBefore=10, spaceAfter=8)
    S_BODY = ParagraphStyle(name='Body', parent=styles['BodyText'], leading=13, spaceAfter=6)
    S_CAP  = ParagraphStyle(name='Cap', parent=styles['BodyText'], alignment=1, fontName='Times-Italic', fontSize=9, spaceAfter=6)
    S_CLAIM= ParagraphStyle(name='Claim', parent=styles['BodyText'], leftIndent=18, firstLineIndent=-18, spaceBefore=6, leading=13)

    BRIEF = [
        "FIG. 1 illustrates a three-tier NBMF memory governed by eDNA with promotion/eviction routing.",
        "FIG. 2 shows tier thresholds and validation for safe promotions.",
        "FIG. 3 depicts L2 quarantine with consensus/divergence scoring and outcomes.",
        "FIG. 4 shows dual-mode encoding converging into NBMF bytecode.",
        "FIG. 5 depicts Merkle-notarized lineage for auditable history.",
        "FIG. 6 illustrates Genome, Epigenome, Lineage, and Immune components.",
        "FIG. 7 shows detect–quarantine–rollback defense loop.",
        "FIG. 8 shows dynamic CPU/GPU/TPU routing via a tensor router.",
        "FIG. 9 shows cross-tenant isolation with sanitized artifacts."
    ]
    def center(drawing, width=6.25*inch):
        t = Table([[drawing]], colWidths=[width])
        t.setStyle(TableStyle([("ALIGN", (0,0), (-1,-1), "CENTER")]))
        return t

    doc = SimpleDocTemplate(str(out_pdf), pagesize=LETTER,
                            leftMargin=0.75*inch, rightMargin=0.75*inch,
                            topMargin=0.8*inch, bottomMargin=0.8*inch)
    story = []
    story.append(Paragraph("<b>PROVISIONAL PATENT APPLICATION</b>", S_TITLE))
    story.append(Paragraph("<b>Title:</b> Neural-Backed Memory Fabric with Enterprise Digital DNA (NBMF-eDNA)", S_H1))
    story.append(Paragraph("<b>Inventor:</b> Masoud Masoori — Richmond Hill, Ontario, Canada", S_BODY))
    story.append(Spacer(1, 10))

    story.append(Paragraph("<b>Brief Description of the Drawings</b>", S_H1))
    for line in BRIEF: story.append(Paragraph("• " + line, S_BODY))

    if not compact: story.append(PageBreak())
    else: story.append(Spacer(1, 8))

    story.append(Paragraph("<b>Detailed Description</b>", S_H1))
    for i, fn in enumerate(funcs, start=1):
        d = fn()
        story.append(center(d))
        story.append(Paragraph(caps[i-1], S_CAP))
        text = details[i-1] if details and len(details) >= i else BRIEF[i-1]
        story.append(Paragraph(text, S_BODY))
        story.append(Spacer(1, 6))

    story.append(PageBreak())
    story.append(Paragraph("<b>Claims</b>", S_H1))
    primary = [(n,t,b) for (n,t,b) in claims if n<=8]
    depend  = [(n,t,b) for (n,t,b) in claims if n>8]
    for (n,title,body) in primary:
        header = f"{n}. {title}" if title else f"{n}."
        story.append(Paragraph(f"<b>{header}</b> {body}", S_CLAIM))
    if depend:
        story.append(Spacer(1,8))
        story.append(Paragraph("<b>Dependent Claims</b>", S_H1))
        for (n,title,body) in depend:
            header = f"{n}. {title}" if title else f"{n}."
            story.append(Paragraph(f"<b>{header}</b> {body}", S_CLAIM))

    doc.build(story)

def build_docx(out_docx: Path, funcs, caps, details, claims, fig_dir: Path):
    if not HAVE_DOCX:
        log("python-docx not installed; skipping DOCX.")
        return
    try:
        from PIL import Image  # ensure Pillow exists
    except Exception:
        log("Pillow not installed; skipping DOCX (needed for PNG rendering).")
        return

    doc = Document()
    doc.add_heading('PROVISIONAL PATENT APPLICATION', 0)
    p = doc.add_paragraph(); p.add_run('Title: ').bold = True
    p.add_run('Neural-Backed Memory Fabric with Enterprise Digital DNA (NBMF-eDNA)')
    p = doc.add_paragraph(); p.add_run('Inventor: ').bold = True
    p.add_run('Masoud Masoori — Richmond Hill, Ontario, Canada')

    doc.add_heading('Brief Description of the Drawings', level=1)
    BRIEF = [
        "FIG. 1 illustrates a three-tier NBMF memory governed by eDNA with promotion/eviction routing.",
        "FIG. 2 shows tier thresholds and validation for safe promotions.",
        "FIG. 3 depicts L2 quarantine with consensus/divergence scoring and outcomes.",
        "FIG. 4 shows dual-mode encoding converging into NBMF bytecode.",
        "FIG. 5 depicts Merkle-notarized lineage for auditable history.",
        "FIG. 6 illustrates Genome, Epigenome, Lineage, and Immune components.",
        "FIG. 7 shows detect–quarantine–rollback defense loop.",
        "FIG. 8 shows dynamic CPU/GPU/TPU routing via a tensor router.",
        "FIG. 9 shows cross-tenant isolation with sanitized artifacts."
    ]
    for line in BRIEF: doc.add_paragraph(line)

    doc.add_heading('Detailed Description', level=1)
    for i, fn in enumerate(funcs, start=1):
        d = fn()
        png = fig_dir / f"FIG{i:02d}.png"
        # Force PIL backend to avoid rlPyCairo dependency
        renderPM.drawToFile(d, str(png), fmt="PNG", backend="pil")
        doc.add_picture(str(png), width=Inches(6.2))
        cap = doc.add_paragraph(); r = cap.add_run(caps[i-1]); r.italic = True
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        body = details[i-1] if details and len(details)>=i else BRIEF[i-1]
        doc.add_paragraph(body)

    doc.add_heading('Claims', level=1)
    style = doc.styles['Normal']; style.font.size = Pt(11)
    def add_claim(n,title,body):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        para.paragraph_format.first_line_indent = Inches(-0.25)
        run = para.add_run(f"{n}. "); run.bold = True
        if title:
            r2 = para.add_run(title + " — "); r2.bold = True
        para.add_run(body)

    primary = [(n,t,b) for (n,t,b) in claims if n<=8]
    depend  = [(n,t,b) for (n,t,b) in claims if n>8]
    for n,t,b in primary: add_claim(n,t,b)
    if depend:
        doc.add_heading('Dependent Claims', level=2)
        for n,t,b in depend: add_claim(n,t,b)

    doc.save(str(out_docx))

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--fig-module", default="generate_patent_color_v2.py")
    ap.add_argument("--zip", default=None)
    ap.add_argument("--compact", action="store_true", help="Start figures on page 1 (reduce gaps)")
    ap.add_argument("--no-docx", action="store_true", help="Skip DOCX even if python-docx is installed")
    args = ap.parse_args()

    base = Path(__file__).resolve().parent
    fig_path = (base / args.fig_module).resolve()
    funcs, caps, details = load_fig_module(fig_path)
    log(f"Using figures from: {fig_path}")

    # Claims
    claims = []
    if args.zip:
        zp = Path(args.zip).expanduser().resolve()
        if zp.exists():
            claims = claims_from_zip(zp)
    if not claims and (base / "CLAIMS_DRAFT.md").exists():
        claims = parse_claims((base / "CLAIMS_DRAFT.md").read_text(encoding="utf-8", errors="ignore"))
    if not claims:
        claims = [(i,"",f"Placeholder claim {i}.") for i in range(1,21)]
    (base / "CLAIMS_SYNCED_v2.txt").write_text("\n".join([f"{n}. {t} — {b}" for n,t,b in claims]), encoding="utf-8")

    # Outputs
    out_pdf  = base / "NBMF_Patent_Application_COLOR_SYNCED_v2.pdf"
    out_docx = base / "NBMF_Patent_Application_v2.docx"
    fig_dir  = base / "figs_v2"; fig_dir.mkdir(exist_ok=True)

    build_pdf(out_pdf, funcs, caps, details, claims, compact=args.compact)
    if HAVE_DOCX and not args.no_docx:
        build_docx(out_docx, funcs, caps, details, claims, fig_dir)
    else:
        log("Skipping DOCX (python-docx not installed or --no-docx set).")

    log(f"Wrote: {out_pdf}")
    if HAVE_DOCX and not args.no_docx and HAVE_DOCX:
        log(f"Wrote: {out_docx}")

if __name__ == "__main__":
    main()
